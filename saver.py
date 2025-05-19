from abc import ABC, abstractmethod
from docx import Document
from openpyxl import Workbook
class Savable(ABC):
    """Абстрактный класс для объектов, которые можно сохранять."""
    @abstractmethod
    def save(self, filename):
        """Абстрактный метод сохранения в файл."""
        pass

class ResultSaver(Savable):
    """Класс для сохранения результатов расчетов."""
    def __init__(self, result_text):
        self.result_text = result_text

    def save_to_doc(self, filename="lab12.docx"):
        """Сохраняет результат в файл .docx."""
        doc = Document()
        doc.add_paragraph("Результат расчёта:")
        doc.add_paragraph(self.result_text)
        doc.save(filename)

    def save_to_xls(self, filename="lab12.xlsx"):
        """Сохраняет результат в файл .xlsx."""
        wb = Workbook()
        ws = wb.active
        ws.title = "Результат"
        ws['A1'] = "Результат расчёта:"
        ws['A2'] = self.result_text
        wb.save(filename)

    def save(self, filename, file_type="doc"):
        """Сохраняет результат в указанный файл."""
        if file_type == "doc":
            self.save_to_doc(filename)
        elif file_type == "xls":
            self.save_to_xls(filename)
        else:
            raise ValueError("Неподдерживаемый тип файла.")