from docx import Document
from openpyxl import Workbook

def save_to_doc(result):
    doc = Document()
    doc.add_paragraph("Результат расчёта:")
    doc.add_paragraph(result)
    doc.save("lab12.docx")

def save_to_xls(result):
    wb = Workbook()
    ws = wb.active
    ws.title = "Результат"
    ws['A1'] = "Результат расчёта:"
    ws['A2'] = result
    wb.save("lab12.xlsx")